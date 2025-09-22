"""Servicio de email con Resend - Simplificado para MVP."""

import resend
import tempfile
import base64
import re
from typing import Optional, Dict, Any
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from voxcliente.config import settings

# Constantes para tipos MIME
WORD_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class ResendEmailService:
    """Servicio simple para Resend."""
    
    def __init__(self):
        """Inicializar cliente de Resend."""
        resend.api_key = settings.resend_api_key
        self.template_path = Path(__file__).parent.parent / "templates" / "email_template.html"
    
    def send_acta_email(self, email: str, acta_data: Dict[str, Any], filename: str, transcript: str = None) -> bool:
        """
        Enviar acta por email usando Resend con archivos Word adjuntos.
        
        Args:
            email: Email del destinatario
            acta_data: Diccionario con resumen_ejecutivo y acta completa
            filename: Nombre del archivo procesado
            transcript: Transcripción completa de Assembly (opcional)
            
        Returns:
            True si se envió correctamente, False si hubo error
        """
        acta_file_path = None
        transcript_file_path = None
        try:
            # Cargar template HTML
            template_content = self._load_template()
            if not template_content:
                print("Error: No se pudo cargar el template de email")
                return False
            
            # Personalizar template con resumen ejecutivo en el cuerpo
            html_content = self._personalize_template(template_content, acta_data, filename)
            
            # Generar archivo Word del acta
            acta_file_path = self._generate_word_document(acta_data, filename)
            if not acta_file_path:
                print("Error: No se pudo generar el archivo Word del acta")
                return False
            
            # Preparar lista de adjuntos
            attachments = []
            
            # Adjunto 1: Acta de reunión
            attachments.append(self._create_attachment(acta_file_path, self._generate_filename("Acta_Reunion", filename)))
            
            # Adjunto 2: Transcripción completa (si está disponible)
            if transcript:
                transcript_file_path = self._generate_transcript_document(transcript, filename)
                if transcript_file_path:
                    attachments.append(self._create_attachment(transcript_file_path, self._generate_filename("Transcripcion_Completa", filename)))
            
            # Preparar datos del email
            email_data = {
                "from": f"{settings.from_name} <{settings.from_email}>",
                "to": [email],
                "subject": f"Acta de Reunión - {filename}",
                "html": html_content,
                "attachments": attachments,
                "reply_to": settings.reply_to_email
            }
            
            # Enviar email
            response = resend.Emails.send(email_data)
            
            print(f"Email enviado exitosamente con {len(attachments)} adjunto(s)")
            return True
            
        except Exception as e:
            print(f"Error enviando email: {e}")
            return False
        finally:
            # Limpiar archivos temporales
            self._cleanup_temp_files([acta_file_path, transcript_file_path])
    
    def _cleanup_temp_files(self, file_paths: list) -> None:
        """Limpiar archivos temporales."""
        for file_path in file_paths:
            if file_path:
                try:
                    Path(file_path).unlink()
                except Exception as e:
                    print(f"Error eliminando archivo temporal: {e}")
    
    def _generate_filename(self, prefix: str, original_filename: str) -> str:
        """Generar nombre de archivo para adjunto."""
        clean_name = original_filename.replace('.', '_')
        return f"{prefix}_{clean_name}.docx"
    
    def _configure_document_margins(self, doc: Document) -> None:
        """Configurar márgenes estándar para documentos Word."""
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def _create_attachment(self, file_path: str, filename: str) -> Dict[str, str]:
        """Crear adjunto para email desde archivo."""
        with open(file_path, 'rb') as file:
            content = file.read()
            content_base64 = base64.b64encode(content).decode('utf-8')
        
        return {
            "filename": filename,
            "content": content_base64,
            "content_type": WORD_MIME_TYPE
        }
    
    def _load_template(self) -> Optional[str]:
        """Cargar template HTML desde archivo."""
        try:
            return self.template_path.read_text(encoding='utf-8')
        except Exception as e:
            print(f"Error cargando template: {e}")
            return None
    
    def _personalize_template(self, template: str, acta_data: Dict[str, Any], filename: str) -> str:
        """Personalizar template con datos específicos."""
        timestamp = datetime.now().strftime("%d/%m/%Y a las %H:%M")
        
        # Extraer resumen ejecutivo
        resumen = acta_data.get("resumen_ejecutivo", {})
        objetivo = resumen.get("objetivo", "No especificado")
        acuerdos = resumen.get("acuerdos", "No especificados")
        proximos_pasos = resumen.get("proximos_pasos", "No especificados")
        
        # Crear resumen ejecutivo para el cuerpo del email
        resumen_ejecutivo_html = f"""
        <div style="background-color: #f8f9fa; padding: 20px; border-radius: 8px; margin: 20px 0;">
            <h3 style="color: #2c3e50; margin-top: 0;">📋 Resumen Ejecutivo</h3>
            <p><strong>🎯 Objetivo:</strong> {objetivo}</p>
            <p><strong>🤝 Acuerdos:</strong> {acuerdos}</p>
            <p><strong>📅 Próximos Pasos:</strong> {proximos_pasos}</p>
        </div>
        <p style="color: #666; font-size: 14px;">
            <em>Para detalles completos, consulta el acta adjunta.</em>
        </p>
        """
        
        # Acta completa para adjunto (si se necesita en el template)
        acta_completa = acta_data.get("acta", "")
        
        return template.replace("{{ resumen_ejecutivo }}", resumen_ejecutivo_html) \
                     .replace("{{ acta_content }}", acta_completa) \
                     .replace("{{ filename }}", filename) \
                     .replace("{{ timestamp }}", timestamp)
    
    def _generate_word_document(self, acta_data: Dict[str, Any], filename: str) -> Optional[str]:
        """
        Generar archivo Word del acta.
        
        Args:
            acta_data: Diccionario con resumen_ejecutivo y acta completa
            filename: Nombre del archivo original
            
        Returns:
            Ruta del archivo Word generado o None si hay error
        """
        try:
            # Crear documento Word
            doc = Document()
            
            # Configurar márgenes estándar
            self._configure_document_margins(doc)
            
            # Título principal
            title = doc.add_heading('ACTA DE REUNIÓN', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Información del archivo
            doc.add_paragraph(f'Archivo procesado: {filename}')
            doc.add_paragraph(f'Fecha de generación: {datetime.now().strftime("%d/%m/%Y a las %H:%M")}')
            doc.add_paragraph('')  # Línea en blanco
            
            # Resumen ejecutivo
            resumen = acta_data.get("resumen_ejecutivo", {})
            if resumen:
                doc.add_heading('RESUMEN EJECUTIVO', level=1)
                
                objetivo = resumen.get("objetivo", "No especificado")
                doc.add_paragraph(f'🎯 Objetivo: {objetivo}')
                
                acuerdos = resumen.get("acuerdos", "No especificados")
                doc.add_paragraph(f'🤝 Acuerdos: {acuerdos}')
                
                proximos_pasos = resumen.get("proximos_pasos", "No especificados")
                doc.add_paragraph(f'📅 Próximos Pasos: {proximos_pasos}')
                
                doc.add_paragraph('')  # Línea en blanco
            
            # Acta completa
            acta_completa = acta_data.get("acta", "")
            if acta_completa:
                doc.add_heading('ACTA COMPLETA', level=1)
                
                # Procesar el texto del acta (convertir markdown básico a formato Word)
                self._add_formatted_text(doc, acta_completa)
            
            # Agregar branding discreto al final del documento
            doc.add_paragraph('')  # Línea en blanco
            doc.add_paragraph('')  # Línea en blanco
            
            # Branding discreto y profesional
            branding_paragraph = doc.add_paragraph()
            branding_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            branding_run = branding_paragraph.add_run('Generado por VoxCliente – Actas profesionales al instante con IA')
            branding_run.font.size = Inches(0.1)  # Tamaño pequeño
            branding_run.font.color.rgb = None  # Color gris por defecto
            
            # Crear archivo temporal
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            temp_path = temp_file.name
            temp_file.close()
            
            # Guardar documento
            doc.save(temp_path)
            
            return temp_path
            
        except Exception as e:
            print(f"Error generando documento Word: {e}")
            return None
    
    def _add_formatted_text(self, doc: Document, text: str):
        """Agregar texto formateado al documento Word."""
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph('')
                continue
            
            # Detectar títulos (líneas que empiezan con números y **)
            if re.match(r'^\d+\.\s*\*\*.*\*\*', line):
                # Es un título principal
                title_text = re.sub(r'^\d+\.\s*\*\*(.*)\*\*', r'\1', line)
                doc.add_heading(title_text, level=2)
            elif line.startswith('**') and line.endswith('**'):
                # Es un subtítulo
                title_text = line[2:-2]
                doc.add_heading(title_text, level=3)
            elif line.startswith('- '):
                # Es un elemento de lista
                item_text = line[2:]
                p = doc.add_paragraph(item_text, style='List Bullet')
            else:
                # Es texto normal
                doc.add_paragraph(line)
    
    def _generate_transcript_document(self, transcript: str, filename: str) -> Optional[str]:
        """
        Generar documento Word con transcripción completa de Assembly.
        
        Args:
            transcript: Transcripción cruda de Assembly con información de hablantes
            filename: Nombre del archivo original
            
        Returns:
            Ruta del archivo Word generado o None si hay error
        """
        try:
            # Crear documento Word
            doc = Document()
            
            # Configurar márgenes estándar
            self._configure_document_margins(doc)
            
            # Título principal
            title = doc.add_heading('TRANSCRIPCIÓN COMPLETA', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Información del archivo
            doc.add_paragraph(f'Archivo procesado: {filename}')
            doc.add_paragraph(f'Fecha de generación: {datetime.now().strftime("%d/%m/%Y a las %H:%M")}')
            doc.add_paragraph('')  # Línea en blanco
            
            # Nota sobre el contenido
            doc.add_paragraph('Esta transcripción contiene el texto completo tal como fue procesado por AssemblyAI, incluyendo la identificación de hablantes.')
            doc.add_paragraph('')  # Línea en blanco
            
            # Transcripción completa (formato crudo)
            doc.add_heading('TRANSCRIPCIÓN', level=1)
            
            # Agregar transcripción línea por línea manteniendo formato original
            lines = transcript.split('\n')
            for line in lines:
                if line.strip():  # Solo agregar líneas no vacías
                    doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph('')  # Mantener saltos de línea
            
            # Agregar branding discreto al final del documento de transcripción
            doc.add_paragraph('')  # Línea en blanco
            doc.add_paragraph('')  # Línea en blanco
            
            # Branding discreto y profesional
            branding_paragraph = doc.add_paragraph()
            branding_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            branding_run = branding_paragraph.add_run('Generado por VoxCliente – Actas profesionales al instante con IA')
            branding_run.font.size = Inches(0.1)  # Tamaño pequeño
            branding_run.font.color.rgb = None  # Color gris por defecto
            
            # Crear archivo temporal
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            temp_path = temp_file.name
            temp_file.close()
            
            # Guardar documento
            doc.save(temp_path)
            
            return temp_path
            
        except Exception as e:
            print(f"Error generando documento de transcripción: {e}")
            return None


# Instancia global del servicio
resend_email_service = ResendEmailService()
