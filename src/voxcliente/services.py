"""Servicios externos - Simplificado para MVP."""

import assemblyai as aai
import openai
import resend
import json
import re
import tempfile
import base64
import os
from typing import Optional, Dict, Any
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from voxcliente.config import settings

# Constantes para tipos MIME
WORD_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class AssemblyAIService:
    """Servicio simple para AssemblyAI."""
    
    def __init__(self):
        """Inicializar cliente de AssemblyAI."""
        aai.settings.api_key = settings.assemblyai_api_key
        config = aai.TranscriptionConfig(
            speaker_labels=True,
            format_text=True,
            punctuate=True,
            language_code="es",
            )

        config.set_custom_spelling(
            {
                "LAIVE": ["laib", "Laibe"],
                "HORECA": ["Oreka"],
            }
        )

        self.transcriber = aai.Transcriber(config=config)
    
    def transcribe_file(self, file_path: str) -> Optional[Dict[str, Any]]:
        """
        Transcribir archivo local usando AssemblyAI con utterances.
        Basado en la documentación oficial de AssemblyAI.
        
        Args:
            file_path: Ruta del archivo local
            
        Returns:
            Diccionario con transcripción y información de costos o None si hay error
        """
        try:
            # Transcribir archivo local directamente
            transcript = self.transcriber.transcribe(file_path)
            
            # Guardar respuesta de AssemblyAI en archivo local para debugging
            self._save_assemblyai_response(transcript, file_path)
            
            # Verificar si la transcripción fue exitosa
            if transcript.status == aai.TranscriptStatus.error:
                print(f"Error en la transcripción: {transcript.error}")
                return None
            
            # Procesar utterances para obtener texto con información de hablantes
            formatted_text = ""
            if transcript.utterances:
                for utterance in transcript.utterances:
                    formatted_text += f"Hablante {utterance.speaker}: {utterance.text}\n"
                formatted_text = formatted_text.strip()
            else:
                # Fallback al texto completo si no hay utterances
                formatted_text = transcript.text if transcript.text else None
            
            if not formatted_text:
                return None
            
            # Calcular duración y costo de AssemblyAI
            # AssemblyAI cobra $0.27 por hora (Universal) = $0.0045 por minuto
            # Usamos Universal por defecto (incluye speaker labels, punctuation, etc.)
            duration_minutes = transcript.audio_duration / 60 if transcript.audio_duration else 0
            assemblyai_cost = duration_minutes * 0.0045
            
            return {
                'transcript': formatted_text,
                'assemblyai_usage': {
                    'audio_duration_seconds': transcript.audio_duration,
                    'audio_duration_minutes': round(duration_minutes, 2),
                    'confidence': transcript.confidence
                },
                'assemblyai_cost': {
                    'duration_minutes': round(duration_minutes, 2),
                    'cost_usd': round(assemblyai_cost, 6)
                }
            }
            
        except Exception as e:
            print(f"Error en transcripción: {e}")
            return None
    
    def _save_assemblyai_response(self, transcript, file_path: str) -> None:
        """Guardar respuesta de AssemblyAI en archivo local para debugging."""
        try:
            # Crear directorio de logs si no existe
            logs_dir = Path("logs")
            logs_dir.mkdir(exist_ok=True)
            
            # Generar nombre de archivo con timestamp al inicio para ordenamiento cronológico
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{timestamp}_assemblyai_response.txt"
            log_file_path = logs_dir / filename
            
            # Extraer información del transcript
            status = transcript.status
            confidence = getattr(transcript, 'confidence', 'N/A')
            audio_duration = getattr(transcript, 'audio_duration', 'N/A')
            text = getattr(transcript, 'text', 'N/A')
            utterances_count = len(transcript.utterances) if transcript.utterances else 0
            
            # Crear contenido del archivo
            content = f"""
=== RESPUESTA DE ASSEMBLYAI ===
Timestamp: {datetime.now().isoformat()}
Archivo procesado: {Path(file_path).name}
Status: {status}
Confidence: {confidence}
Duración audio (segundos): {audio_duration}
Número de utterances: {utterances_count}

=== TEXTO COMPLETO ===
{text}

=== UTTERANCES (HABLANTES) ===
"""
            
            # Agregar utterances si existen
            if transcript.utterances:
                for i, utterance in enumerate(transcript.utterances):
                    content += f"Utterance {i+1}:\n"
                    content += f"  Speaker: {utterance.speaker}\n"
                    content += f"  Text: {utterance.text}\n"
                    content += f"  Confidence: {getattr(utterance, 'confidence', 'N/A')}\n"
                    content += f"  Start: {getattr(utterance, 'start', 'N/A')}ms\n"
                    content += f"  End: {getattr(utterance, 'end', 'N/A')}ms\n\n"
            else:
                content += "No hay utterances disponibles\n"
            
            content += f"""
=== METADATOS COMPLETOS ===
{transcript}

=== FIN DE RESPUESTA ===
"""
            
            # Guardar archivo
            with open(log_file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            print(f"📁 Respuesta de AssemblyAI guardada en: {log_file_path}")
            
        except Exception as e:
            print(f"⚠️ Error guardando respuesta de AssemblyAI: {e}")


class OpenAIService:
    """Servicio simple para OpenAI."""
    
    def __init__(self):
        """Inicializar cliente de OpenAI."""
        self.client = openai.OpenAI(api_key=settings.openai_api_key)
        self.prompt_path = Path(__file__).parent / "prompts" / "acta_generation.txt"
    
    def generate_acta(self, transcript: str) -> Optional[Dict[str, Any]]:
        """
        Generar acta profesional a partir de transcripción.
        
        Args:
            transcript: Transcripción del audio
            
        Returns:
            Diccionario con resumen_ejecutivo, acta completa y costos o None si hay error
        """
        try:
            # Cargar prompt desde archivo
            prompt_template = self._load_prompt()
            if not prompt_template:
                print("Error: No se pudo cargar el prompt")
                return None
            
            response = self.client.chat.completions.create(
                model="gpt-5-mini",
                messages=[
                    {"role": "system", "content": prompt_template},
                    {"role": "user", "content": transcript}
                ],
                # temperature=0.3 # gpt-5-mini no soporta temperature
            )
            
            # Parsear respuesta JSON
            raw_response = response.choices[0].message.content
            
            # Guardar respuesta de OpenAI en archivo local para debugging
            self._save_openai_response(raw_response, transcript[:50])
            
            parsed_data = self._parse_response(raw_response)
            
            # Agregar información de costos
            if parsed_data:
                parsed_data['openai_usage'] = {
                    'prompt_tokens': response.usage.prompt_tokens,
                    'completion_tokens': response.usage.completion_tokens,
                    'total_tokens': response.usage.total_tokens
                }
                
                # Calcular costo real (precios de GPT-5-mini)
                # Input: $0.25 per 1M tokens = $0.00025 per 1K tokens
                # Output: $2 per 1M tokens = $0.002 per 1K tokens
                input_cost = (response.usage.prompt_tokens / 1000) * 0.00025
                output_cost = (response.usage.completion_tokens / 1000) * 0.002
                total_openai_cost = input_cost + output_cost
                
                parsed_data['openai_cost'] = {
                    'input_cost_usd': round(input_cost, 6),
                    'output_cost_usd': round(output_cost, 6),
                    'total_cost_usd': round(total_openai_cost, 6)
                }
            
            return parsed_data
            
        except Exception as e:
            print(f"Error generando acta: {e}")
            return None
    
    def _load_prompt(self) -> Optional[str]:
        """Cargar prompt desde archivo."""
        try:
            return self.prompt_path.read_text(encoding='utf-8')
        except Exception as e:
            print(f"Error cargando prompt: {e}")
            return None
    
    def _save_openai_response(self, response: str, transcript_preview: str) -> None:
        """Guardar respuesta de OpenAI en archivo local para debugging."""
        try:
            # Crear directorio de logs si no existe
            logs_dir = Path("logs")
            logs_dir.mkdir(exist_ok=True)
            
            # Generar nombre de archivo con timestamp al inicio para ordenamiento cronológico
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{timestamp}_openai_response.txt"
            file_path = logs_dir / filename
            
            # Crear contenido del archivo
            content = f"""
=== RESPUESTA DE OPENAI ===
Timestamp: {datetime.now().isoformat()}
Modelo: gpt-5-mini
Transcripción (primeros 50 chars): {transcript_preview}

=== RESPUESTA COMPLETA ===
{response}

=== FIN DE RESPUESTA ===
"""
            
            # Guardar archivo
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            print(f"📁 Respuesta de OpenAI guardada en: {file_path}")
            
        except Exception as e:
            print(f"⚠️ Error guardando respuesta de OpenAI: {e}")
    
    def _parse_response(self, raw_response: str) -> Optional[Dict[str, Any]]:
        """Parsear respuesta JSON de OpenAI."""
        try:
            # Intentar múltiples formatos de parsing
            
            # 1. Buscar JSON entre etiquetas <output> y </output>
            pattern = r'<output>(.*?)</output>'
            match = re.search(pattern, raw_response, re.DOTALL)
            
            if match:
                json_str = match.group(1).strip()
                print(f"✅ JSON encontrado entre etiquetas <output>: {json_str[:100]}...")
            else:
                # 2. Buscar JSON entre ```json y ```
                pattern = r'```json\s*(.*?)\s*```'
                match = re.search(pattern, raw_response, re.DOTALL)
                
                if match:
                    json_str = match.group(1).strip()
                    print(f"✅ JSON encontrado entre ```json: {json_str[:100]}...")
                else:
                    # 3. Buscar JSON directo (sin etiquetas)
                    json_str = raw_response.strip()
                    print(f"⚠️ Intentando parsear respuesta directa: {json_str[:100]}...")
            
            # Parsear JSON
            parsed_data = json.loads(json_str)
            
            # Validar estructura requerida
            if not self._validate_structure(parsed_data):
                print("Error: Estructura JSON inválida")
                return None
            
            return parsed_data
            
        except json.JSONDecodeError as e:
            print(f"❌ Error parseando JSON: {e}")
            print(f"📄 Respuesta completa de OpenAI: {raw_response}")
            return None
        except Exception as e:
            print(f"❌ Error procesando respuesta: {e}")
            print(f"📄 Respuesta completa de OpenAI: {raw_response}")
            return None
    
    def _validate_structure(self, data: Dict[str, Any]) -> bool:
        """Validar estructura del JSON parseado."""
        try:
            # Verificar que tenga las claves principales
            if "resumen_ejecutivo" not in data or "acta" not in data:
                return False
            
            resumen = data["resumen_ejecutivo"]
            if not isinstance(resumen, dict):
                return False
            
            # Verificar campos del resumen ejecutivo
            required_fields = ["objetivo", "acuerdos", "proximos_pasos"]
            for field in required_fields:
                if field not in resumen:
                    return False
            
            return True
            
        except Exception:
            return False


class ResendEmailService:
    """Servicio simple para Resend."""
    
    def __init__(self):
        """Inicializar cliente de Resend."""
        resend.api_key = settings.resend_api_key
        self.template_path = Path(__file__).parent / "templates" / "email_template.html"
    
    def send_acta_email(self, email: str, acta_data: Dict[str, Any], filename: str, transcript: str = None) -> bool:
        """
        Enviar acta por email usando Resend con archivos Word adjuntos.
        
        Args:
            email: Email del destinatario
            acta_data: Diccionario con resumen_ejecutivo y acta completa
            filename: Nombre del archivo procesado
            transcript: Transcripción completa de Assembly (opcional)
            
        Returns:
            True si se envió correctamente, False si hubo error
        """
        acta_file_path = None
        transcript_file_path = None
        try:
            # Cargar template HTML
            template_content = self._load_template()
            if not template_content:
                print("Error: No se pudo cargar el template de email")
                return False
            
            # Personalizar template con resumen ejecutivo en el cuerpo
            html_content = self._personalize_template(template_content, acta_data, filename)
            
            # Generar archivo Word del acta
            acta_file_path = self._generate_word_document(acta_data, filename)
            if not acta_file_path:
                print("Error: No se pudo generar el archivo Word del acta")
                return False
            
            # Preparar lista de adjuntos
            attachments = []
            
            # Adjunto 1: Acta de reunión
            attachments.append(self._create_attachment(acta_file_path, self._generate_filename("Acta_Reunion", filename)))
            
            # Adjunto 2: Transcripción completa (si está disponible)
            if transcript:
                transcript_file_path = self._generate_transcript_document(transcript, filename)
                if transcript_file_path:
                    attachments.append(self._create_attachment(transcript_file_path, self._generate_filename("Transcripcion_Completa", filename)))
            
            # Preparar datos del email
            email_data = {
                "from": f"{settings.from_name} <{settings.from_email}>",
                "to": [email],
                "subject": f"Acta de Reunión - {filename}",
                "html": html_content,
                "attachments": attachments
            }
            
            # Enviar email
            response = resend.Emails.send(email_data)
            
            print(f"Email enviado exitosamente con {len(attachments)} adjunto(s)")
            return True
            
        except Exception as e:
            print(f"Error enviando email: {e}")
            return False
    
    def _cleanup_temp_files(self, file_paths: list) -> None:
        """Limpiar archivos temporales."""
        for file_path in file_paths:
            if file_path:
                try:
                    Path(file_path).unlink()
                except Exception as e:
                    print(f"Error eliminando archivo temporal: {e}")
    
    def _generate_filename(self, prefix: str, original_filename: str) -> str:
        """Generar nombre de archivo para adjunto."""
        clean_name = original_filename.replace('.', '_')
        return f"{prefix}_{clean_name}.docx"
    
    def _configure_document_margins(self, doc: Document) -> None:
        """Configurar márgenes estándar para documentos Word."""
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def _create_attachment(self, file_path: str, filename: str) -> Dict[str, str]:
        """Crear adjunto para email desde archivo."""
        with open(file_path, 'rb') as file:
            content = file.read()
            content_base64 = base64.b64encode(content).decode('utf-8')
        
        return {
            "filename": filename,
            "content": content_base64,
            "content_type": WORD_MIME_TYPE
        }
    
    def _load_template(self) -> Optional[str]:
        """Cargar template HTML desde archivo."""
        try:
            return self.template_path.read_text(encoding='utf-8')
        except Exception as e:
            print(f"Error cargando template: {e}")
            return None
    
    def _personalize_template(self, template: str, acta_data: Dict[str, Any], filename: str) -> str:
        """Personalizar template con datos específicos."""
        timestamp = datetime.now().strftime("%d/%m/%Y a las %H:%M")
        
        # Extraer resumen ejecutivo
        resumen = acta_data.get("resumen_ejecutivo", {})
        objetivo = resumen.get("objetivo", "No especificado")
        acuerdos = resumen.get("acuerdos", "No especificados")
        proximos_pasos = resumen.get("proximos_pasos", "No especificados")
        
        # Crear resumen ejecutivo para el cuerpo del email
        resumen_ejecutivo_html = f"""
        <div style="background-color: #f8f9fa; padding: 20px; border-radius: 8px; margin: 20px 0;">
            <h3 style="color: #2c3e50; margin-top: 0;">📋 Resumen Ejecutivo</h3>
            <p><strong>🎯 Objetivo:</strong> {objetivo}</p>
            <p><strong>🤝 Acuerdos:</strong> {acuerdos}</p>
            <p><strong>📅 Próximos Pasos:</strong> {proximos_pasos}</p>
        </div>
        <p style="color: #666; font-size: 14px;">
            <em>Para detalles completos, consulta el acta adjunta.</em>
        </p>
        """
        
        # Acta completa para adjunto (si se necesita en el template)
        acta_completa = acta_data.get("acta", "")
        
        return template.replace("{{ resumen_ejecutivo }}", resumen_ejecutivo_html) \
                     .replace("{{ acta_content }}", acta_completa) \
                     .replace("{{ filename }}", filename) \
                     .replace("{{ timestamp }}", timestamp)
    
    def _generate_word_document(self, acta_data: Dict[str, Any], filename: str) -> Optional[str]:
        """
        Generar archivo Word del acta.
        
        Args:
            acta_data: Diccionario con resumen_ejecutivo y acta completa
            filename: Nombre del archivo original
            
        Returns:
            Ruta del archivo Word generado o None si hay error
        """
        try:
            # Crear documento Word
            doc = Document()
            
            # Configurar márgenes estándar
            self._configure_document_margins(doc)
            
            # Título principal
            title = doc.add_heading('ACTA DE REUNIÓN', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Información del archivo
            doc.add_paragraph(f'Archivo procesado: {filename}')
            doc.add_paragraph(f'Fecha de generación: {datetime.now().strftime("%d/%m/%Y a las %H:%M")}')
            doc.add_paragraph('')  # Línea en blanco
            
            # Resumen ejecutivo
            resumen = acta_data.get("resumen_ejecutivo", {})
            if resumen:
                doc.add_heading('RESUMEN EJECUTIVO', level=1)
                
                objetivo = resumen.get("objetivo", "No especificado")
                doc.add_paragraph(f'🎯 Objetivo: {objetivo}')
                
                acuerdos = resumen.get("acuerdos", "No especificados")
                doc.add_paragraph(f'🤝 Acuerdos: {acuerdos}')
                
                proximos_pasos = resumen.get("proximos_pasos", "No especificados")
                doc.add_paragraph(f'📅 Próximos Pasos: {proximos_pasos}')
                
                doc.add_paragraph('')  # Línea en blanco
            
            # Acta completa
            acta_completa = acta_data.get("acta", "")
            if acta_completa:
                doc.add_heading('ACTA COMPLETA', level=1)
                
                # Procesar el texto del acta (convertir markdown básico a formato Word)
                self._add_formatted_text(doc, acta_completa)
            
            # Crear archivo temporal
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            temp_path = temp_file.name
            temp_file.close()
            
            # Guardar documento
            doc.save(temp_path)
            
            return temp_path
            
        except Exception as e:
            print(f"Error generando documento Word: {e}")
            return None
    
    def _add_formatted_text(self, doc: Document, text: str):
        """Agregar texto formateado al documento Word."""
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph('')
                continue
            
            # Detectar títulos (líneas que empiezan con números y **)
            if re.match(r'^\d+\.\s*\*\*.*\*\*', line):
                # Es un título principal
                title_text = re.sub(r'^\d+\.\s*\*\*(.*)\*\*', r'\1', line)
                doc.add_heading(title_text, level=2)
            elif line.startswith('**') and line.endswith('**'):
                # Es un subtítulo
                title_text = line[2:-2]
                doc.add_heading(title_text, level=3)
            elif line.startswith('- '):
                # Es un elemento de lista
                item_text = line[2:]
                p = doc.add_paragraph(item_text, style='List Bullet')
            else:
                # Es texto normal
                doc.add_paragraph(line)
    
    def _generate_transcript_document(self, transcript: str, filename: str) -> Optional[str]:
        """
        Generar documento Word con transcripción completa de Assembly.
        
        Args:
            transcript: Transcripción cruda de Assembly con información de hablantes
            filename: Nombre del archivo original
            
        Returns:
            Ruta del archivo Word generado o None si hay error
        """
        try:
            # Crear documento Word
            doc = Document()
            
            # Configurar márgenes estándar
            self._configure_document_margins(doc)
            
            # Título principal
            title = doc.add_heading('TRANSCRIPCIÓN COMPLETA', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Información del archivo
            doc.add_paragraph(f'Archivo procesado: {filename}')
            doc.add_paragraph(f'Fecha de generación: {datetime.now().strftime("%d/%m/%Y a las %H:%M")}')
            doc.add_paragraph('')  # Línea en blanco
            
            # Nota sobre el contenido
            doc.add_paragraph('Esta transcripción contiene el texto completo tal como fue procesado por AssemblyAI, incluyendo la identificación de hablantes.')
            doc.add_paragraph('')  # Línea en blanco
            
            # Transcripción completa (formato crudo)
            doc.add_heading('TRANSCRIPCIÓN', level=1)
            
            # Agregar transcripción línea por línea manteniendo formato original
            lines = transcript.split('\n')
            for line in lines:
                if line.strip():  # Solo agregar líneas no vacías
                    doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph('')  # Mantener saltos de línea
            
            # Crear archivo temporal
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            temp_path = temp_file.name
            temp_file.close()
            
            # Guardar documento
            doc.save(temp_path)
            
            return temp_path
            
        except Exception as e:
            print(f"Error generando documento de transcripción: {e}")
            return None


# Instancias globales de los servicios
assemblyai_service = AssemblyAIService()
openai_service = OpenAIService()
resend_email_service = ResendEmailService()
